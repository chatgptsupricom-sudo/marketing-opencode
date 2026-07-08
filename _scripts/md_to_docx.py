import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_table_with_shading(doc, rows_data):
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                paragraph.space_after = Pt(2)
                paragraph.space_before = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A1A2E"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()
    return table

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)

# ════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BASE DE DATOS')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('TEMAS GENÉRICOS B2B SUPRICOM')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xCC, 0x7A, 0x00)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('Aplicable a cualquier marca, producto o categoría del catálogo')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('Julio 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ════════════════════════════════════════════
# NOTA TÉCNICA
# ════════════════════════════════════════════
add_heading_styled(doc, 'NOTA TÉCNICA', 1)
add_heading_styled(doc, 'El buyer B2B tech venezolano (psicografía validada)', 2)

doc.add_paragraph(
    'Perfil: Dueño/gerente de tienda de tecnología, revendedor, integrador. '
    '30–55 años. Empresa de 2–15 empleados. Opera en Venezuela.'
)

add_heading_styled(doc, 'Factores decisivos de compra (prioridad)', 3)
factores = [
    'Disponibilidad inmediata de producto (stock hoy)',
    'Precio competitivo visible (sin tener que pedir cotización)',
    'Facilidad de comunicación (WhatsApp, respuesta rápida)',
    'Garantía real y soporte (productos originales)',
]
for f in factores:
    doc.add_paragraph(f, style='List Number')

add_heading_styled(doc, 'Comportamiento de consumo de contenido', 3)
comportamientos = [
    ('78%', 'prefiere cotizar por WhatsApp (no email, no formularios)'),
    ('—', 'Compara con 2–3 proveedores antes de decidir'),
    ('—', 'Timeline de decisión: 3–7 días (no meses como B2B corporativo)'),
    ('73%', 'del proceso lo investiga en silencio sin contactar al proveedor'),
    ('29%', 'ya inicia investigación en ChatGPT/IA generativa'),
    ('94%', 'ya tiene proveedor favorito en mente antes de contactar'),
    ('⚠️', 'Odian el spam de WhatsApp con listas de precios sin contexto'),
]
for pct, desc in comportamientos:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{pct} ')
    run.bold = True
    p.add_run(desc)

add_heading_styled(doc, 'Lo que funciona con este perfil', 3)
funciona = [
    'Contenido educativo útil (4x más engagement que contenido promocional)',
    'Precios visibles desde el anuncio (3x mejor conversión)',
    'Carruseles > imagen estática > video largo',
    'Respuesta en <5 minutos en WhatsApp (2.8x más conversión)',
    'Contenido que les enseñe a venderle MEJOR a SUS clientes',
    'Evidencias de rotación: "esto se está vendiendo"',
]
for f in funciona:
    doc.add_paragraph(f, style='List Bullet')

add_heading_styled(doc, 'Lo que NO funciona', 3)
nofunciona = [
    'Spam semanal de lista de precios',
    'Videos de 30+ segundos (no los terminan)',
    'Mensajes automatizados largos en WhatsApp',
    '"Cotiza por WhatsApp" genérico sin precio visible',
    'Contenido demasiado técnico (hablan con dueños de tienda, no con ingenieros)',
]
for n in nofunciona:
    doc.add_paragraph(n, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════════
# ARQUETIPOS
# ════════════════════════════════════════════
add_heading_styled(doc, 'ARQUETIPOS DE CONTENIDO', 1)

arquetipos = [
    ('ARQUETIPO 1: PRUEBA SOCIAL / "LO QUE SE VENDE"',
     'Fórmula: validación externa + tendencia + llamado a la acción',
     [
         ('1', 'Estas son las [categoría/marca] más vendidas de [mes/año]', 'Cualquier categoría', 'Carrusel 8 slides'),
         ('2', 'El top 5 de [producto] que tus clientes están pidiendo esta semana', 'Cualquier producto', 'Post + carrusel'),
         ('3', 'Lo que otros dueños de tienda están comprando en [categoría]', 'Cualquier categoría', 'Carrusel'),
         ('4', '[Marca] revela sus 3 productos estrella del trimestre', 'Marcas con portafolio', 'Post + carrusel'),
         ('5', 'Los [producto] con mayor rotación en [mes]', 'Cualquier producto', 'Post + reel 15s'),
         ('6', 'Esto es lo que más se vende en [categoría] según datos de SUPRICOM', 'Cualquier categoría', 'Carrusel 10 slides'),
         ('7', '¿Sabes cuál es el [producto] más buscado de la semana?', 'Cualquier producto', 'Post + Stories'),
         ('8', 'Nuestros clientes mayoristas están llevando [X] unidades de [producto] esta semana', 'Producto con buen stock', 'Post'),
         ('9', 'Las [categoría] que nadie se esperaba que fueran tendencia este mes', 'Categorías nicho', 'Carrusel'),
         ('10', '[Número] de tiendas ya tienen [producto/marca] en su vitrina. ¿Tú aún no?', 'Marcas nuevas/destacadas', 'Post'),
     ]),
    ('ARQUETIPO 2: CASO DE NEGOCIO / "POR QUÉ TENERLO"',
     'Fórmula: problema del cliente final + solución que vendes + beneficio para la tienda',
     [
         ('11', 'Por qué tener [marca/producto] en tu vitrina', 'Cualquier producto', 'Carrusel 10 slides'),
         ('12', '3 razones para sumar [categoría] a tu catálogo ahora', 'Cualquier categoría', 'Post + carrusel'),
         ('13', 'El margen que estás perdiendo si no vendes [categoría]', 'Categorías subestimadas', 'Carrusel'),
         ('14', 'Por qué tus clientes te van a pedir [producto] este mes', 'Estacional', 'Post'),
         ('15', '[Producto] no es un lujo: es una necesidad para tus clientes', 'Protección, respaldo', 'Carrusel'),
         ('16', 'La categoría [X] creció [%] este año. ¿Ya la tienes en tu oferta?', 'Categorías en crecimiento', 'Carrusel + reel'),
         ('17', 'Si no vendes [categoría], le estás regalando ventas a tu competencia', 'Categorías con alta demanda', 'Post + carrusel'),
         ('18', '[Producto]: el artículo que todo dueño de tienda debería tener en inventario', 'Productos básicos', 'Post'),
         ('19', 'Lo que [marca] significa para tu negocio: stock, garantía y rotación', 'Marcas principales', 'Carrusel'),
         ('20', '¿Todavía no tienes [categoría] en tu catálogo? Esto es lo que te estás perdiendo', 'Categorías nuevas', 'Carrusel 8 slides'),
     ]),
    ('ARQUETIPO 3: EDUCACIÓN AL REVENDEDOR / "CÓMO VENDERLO"',
     'Fórmula: objeción del consumidor final + argumento de venta + cierre',
     [
         ('21', 'Cómo vender [producto] sin que el cliente sienta que es caro', 'Productos de precio alto', 'Carrusel 10 slides'),
         ('22', '3 argumentos de venta que cierran más rápido con [producto]', 'Cualquier producto', 'Post + carrusel'),
         ('23', 'Lo que tu cliente NO sabe sobre [producto] y deberías explicarle', 'Productos técnicos', 'Carrusel'),
         ('24', 'Las 5 objeciones más comunes al vender [producto] (y cómo resolverlas)', 'Cualquier producto', 'Carrusel 12 slides'),
         ('25', 'Cómo aumentar el ticket promedio con [categoría]', 'Categorías complementarias', 'Post + carrusel'),
         ('26', 'El discurso de venta de 30 segundos para [producto]', 'Cualquier producto', 'Reel 30s'),
         ('27', '¿Tu cliente duda en comprar [producto]? Dile esto', 'Cualquier producto', 'Post'),
         ('28', 'Cómo presentar [marca] frente a otras opciones del mercado', 'Marcas comparables', 'Carrusel'),
         ('29', '3 formas de hacer upsell con [categoría]', 'Categorías con gamas', 'Carrusel'),
         ('30', 'El error más común al vender [producto] (y cómo evitarlo)', 'Cualquier producto', 'Post + reel'),
     ]),
    ('ARQUETIPO 4: INTELIGENCIA DE MERCADO / "TENDENCIAS"',
     'Fórmula: dato del mercado + implicación para la tienda + acción recomendada',
     [
         ('36', 'Por qué [mes] es el mejor momento para vender [categoría]', 'Categorías estacionales', 'Carrusel'),
         ('37', '[Tendencia del mercado] está impulsando la demanda de [categoría]', 'Categorías vinculadas a crisis', 'Post + carrusel'),
         ('38', 'Lo que la [situación actual] significa para tu negocio de [categoría]', 'Contexto país', 'Carrusel'),
         ('39', '[Categoría] vs [categoría]: por qué una está creciendo más que la otra', 'Categorías comparables', 'Carrusel 10 slides'),
         ('40', 'Datos que todo revendedor de [categoría] debería conocer en 2026', 'Cualquier categoría', 'Carrusel'),
         ('41', 'El mercado de [categoría] en Venezuela: números, tendencias y oportunidades', 'Categorías grandes', 'Carrusel 12 slides'),
         ('42', '¿Hacia dónde va la demanda de [categoría] en los próximos meses?', 'Proyección', 'Post + carrusel'),
         ('43', 'Lo que los consumidores finales están buscando en [categoría] (datos 2026)', 'Cualquier categoría', 'Carrusel'),
         ('44', '3 señales de que [categoría] será la más vendida del próximo trimestre', 'Categorías emergentes', 'Post + carrusel'),
         ('45', 'Así está cambiando el consumo de [categoría] en Venezuela', 'Categorías con evolución', 'Carrusel'),
     ]),
    ('ARQUETIPO 5: OPERACIONAL / "STOCK Y DISPONIBILIDAD"',
     'Fórmula: disponibilidad + urgencia + beneficio de comprar ahora',
     [
         ('46', 'Tenemos [X] unidades de [producto] listas para entregar hoy', 'Producto con stock alto', 'Post'),
         ('47', 'Stock actualizado de [categoría]: esto es lo que hay disponible ahora', 'Categorías', 'Post + carrusel'),
         ('48', 'Acaba de llegar [producto/marca] a SUPRICOM. Stock limitado', 'Nuevos ingresos', 'Post'),
         ('49', 'Los [producto] que se están agotando más rápido esta semana', 'Productos con alta rotación', 'Post + Stories'),
         ('50', 'Así está nuestro inventario de [categoría] hoy', 'Categorías', 'Reel 15s + post'),
         ('51', '[Producto]: disponible hoy con entrega inmediata', 'Cualquier producto', 'Post'),
         ('52', 'Repusimos stock de [producto/marca] — lo que nos pidieron vuelve a estar disponible', 'Marcas con re-stock', 'Post'),
         ('53', '¿Necesitas [categoría] para esta semana? Esto tenemos en almacén', 'Categorías', 'Post + carrusel'),
         ('54', 'La lista de [producto] que sí tenemos en stock', 'Categorías con variantes', 'Post'),
         ('55', 'Pedido especial: [producto] con [X] días de entrega', 'Productos no en stock base', 'Post'),
     ]),
    ('ARQUETIPO 6: COMPARATIVAS / "VS"',
     'Fórmula: dos opciones + criterios + recomendación basada en el caso de uso',
     [
         ('56', '[Producto A] vs [Producto B]: ¿cuál conviene más a tu cliente?', 'Dos productos similares', 'Carrusel 10 slides'),
         ('57', '[Marca A] vs [Marca B]: lo que necesitas saber para recomendar', 'Marcas competidoras', 'Carrusel'),
         ('58', '[Producto gama entrada] vs [Producto gama alta]: ¿cuándo vender cada uno?', 'Gamas del mismo producto', 'Carrusel'),
         ('59', '3 diferencias clave entre [producto] y [alternativa] que tus clientes preguntarán', 'Productos confundibles', 'Post + carrusel'),
         ('60', '¿[Producto] o [alternativa]? Guía rápida para tu cliente indeciso', 'Decisiones comunes', 'Carrusel 8 slides'),
         ('61', 'No es lo mismo [producto A] que [producto B]: esto cambia', 'Productos similares', 'Post'),
         ('62', '¿Vale la pena pagar más por [producto premium]? Te decimos cuándo sí', 'Gama alta vs estándar', 'Carrusel'),
         ('63', 'Comparativa rápida: [marca] vs otras opciones del mercado', 'Marca vs competencia', 'Carrusel'),
         ('64', 'Lo que [marca] tiene que las otras no', 'Ventaja competitiva', 'Post + carrusel'),
         ('65', '¿Cuál [producto] elegir según el tipo de cliente? [3 perfiles]', 'Segmentación', 'Carrusel'),
     ]),
    ('ARQUETIPO 7: ESTACIONAL / "TIMING"',
     'Fórmula: evento/fecha + necesidad asociada + producto como solución',
     [
         ('66', 'Prepárate para [evento/estación]: lo que más se vende en esta época', 'Temporada', 'Carrusel'),
         ('67', '[Mes]: el mes de mayor demanda de [categoría]. ¿Tienes stock?', 'Categorías estacionales', 'Post'),
         ('68', 'Cuando empiezan las lluvias, lo que más necesitan tus clientes es [categoría]', 'Estacional climática', 'Carrusel'),
         ('69', 'Vuelta a clases: [categoría] que no puede faltar en tu vitrina', 'Evento recurrente', 'Carrusel'),
         ('70', 'Black Friday / [evento]: la categoría que más se vende', 'Eventos comerciales', 'Post + carrusel'),
         ('71', 'Julio es el mes de [categoría]: por qué y cómo prepararte', 'Mes temático', 'Carrusel'),
         ('72', '[Época] + [situación país] = demanda récord de [categoría]', 'Coyuntura', 'Post'),
         ('73', 'Lo que aprendimos: [categoría] se vendió [%] más la temporada pasada', 'Histórico', 'Carrusel'),
         ('74', '[Evento]: la oportunidad que no puedes dejar pasar para vender [categoría]', 'Fechas clave', 'Post'),
         ('75', 'Calendario: qué [categoría] vender cada mes del año', 'Planificación anual', 'Carrusel 12 slides'),
     ]),
    ('ARQUETIPO 8: CONFIANZA / "DETRÁS DEL MAYORISTA"',
     'Fórmula: transparencia + proceso + valor agregado',
     [
         ('76', 'Así trabajamos: nuestro proceso de selección de [categoría]', 'Cualquier categoría', 'Reel 30s + carrusel'),
         ('77', '¿Cómo elegimos qué [producto/marca] traer a nuestro catálogo?', 'Transparencia', 'Post + carrusel'),
         ('78', 'El equipo que hace posible tu negocio', 'Detrás de escena', 'Reel'),
         ('79', 'Por qué [marca] confía en SUPRICOM como distribuidor autorizado', 'Marcas representadas', 'Post'),
         ('80', 'Garantía y soporte: qué pasa si tu cliente tiene un problema con [producto]', 'Postventa', 'Carrusel'),
         ('81', 'Nuestra relación con [marca]: lo que significa ser distribuidor oficial', 'Marcas principales', 'Post'),
         ('82', '¿Cómo llega [producto] desde el fabricante hasta tu tienda?', 'Logística', 'Carrusel + reel'),
         ('83', 'Lo que nos diferencia como mayorista de [categoría]', 'Categorías con competencia', 'Carrusel'),
         ('84', '[Años] distribuyendo [categoría] en Venezuela: lo que hemos aprendido', 'Trayectoria', 'Post + carrusel'),
         ('85', 'Los controles de calidad por los que pasa cada [producto] antes de llegar a ti', 'Procesos', 'Reel 30s'),
     ]),
    ('ARQUETIPO 9: RENTABILIDAD / "MARGEN Y GANANCIA"',
     'Fórmula: inversión + retorno + proyección para la tienda',
     [
         ('86', '¿Cuánto puedes ganar por unidad vendida de [producto]?', 'Productos con buen margen', 'Carrusel'),
         ('87', 'El ticket promedio más alto está en [categoría]', 'Categorías alto valor', 'Post + carrusel'),
         ('88', '[Producto]: margen [X]%, rotación [Y] — la cuenta que todo dueño debe hacer', 'Producto estrella', 'Carrusel'),
         ('89', 'Con [X] ventas de [producto] al mes recuperas tu inversión', 'Productos precio medio', 'Post'),
         ('90', 'Los [producto] con mejor relación margen-rotación en nuestro catálogo', 'Optimización', 'Carrusel 10 slides'),
         ('91', '¿Sabes cuánto deja [categoría] al mes? Esto dice el mercado', 'Por categoría', 'Carrusel'),
         ('92', 'Inversión inicial vs retorno mensual: el caso de [categoría]', 'Categorías nuevas', 'Post + carrusel'),
         ('93', '[Producto]: bajo precio, alta rotación, margen constante', 'Productos commodity', 'Post'),
         ('94', 'Por qué tener [categoría] en tu mix mejora tu margen promedio', 'Mix estratégico', 'Carrusel'),
         ('95', 'El TOP 10 de productos con mejor rentabilidad para tu tienda', 'Multi-categoría', 'Carrusel 12 slides'),
     ]),
    ('ARQUETIPO 10: VENTA CRUZADA / "KITS Y COMBOS"',
     'Fórmula: producto ancla + complemento + ahorro percibido',
     [
         ('96', 'El kit perfecto de [categoría] para empezar a vender', 'Starter pack', 'Carrusel'),
         ('97', '3 combos de [categoría] que puedes armar hoy y vender como paquete', 'Complementarias', 'Carrusel'),
         ('98', 'Lo que todo cliente de [producto] termina comprando después', 'Secuencia', 'Post + carrusel'),
         ('99', '[Producto A] + [Producto B]: la combinación que más sale este mes', 'Cross-sell', 'Post'),
         ('100', 'Arma tu propio kit de [categoría] y aumenta el ticket promedio', 'Personalizable', 'Carrusel'),
         ('101', '5 accesorios de [categoría] que deberías vender junto a [producto principal]', 'Complementos', 'Carrusel 8 slides'),
         ('102', 'La canasta básica de [categoría]: lo mínimo que debes tener', 'Inventario mínimo', 'Carrusel'),
         ('103', '¿Vendes [producto] solo? Tus clientes también necesitan [complemento]', 'Cross-sell lógico', 'Post'),
         ('104', 'El combo que todo [tipo de cliente final] necesita', 'Segmento específico', 'Carrusel'),
         ('105', 'De [básico] a [profesional]: la escalera de compra', 'Gamas', 'Carrusel'),
     ]),
    ('ARQUETIPO 11: OBJECIONES DEL DUEÑO DE TIENDA',
     'Fórmula: objeción que él/ella tiene como comprador + respuesta + datos',
     [
         ('106', '"Es muy caro": cómo vender [producto] cuando el precio parece alto', 'Premium', 'Carrusel'),
         ('107', '"¿Y si no se vende?": cómo minimizar el riesgo al probar [categoría] nueva', 'Nuevas', 'Post + carrusel'),
         ('108', '"Es lo mismo que [alternativa]": el argumento para defender [marca]', 'Marcas', 'Carrusel'),
         ('109', '"Mis clientes no piden eso": cómo crear demanda de [categoría]', 'Desconocidas', 'Post + carrusel'),
         ('110', '"No tengo espacio en la vitrina": cómo exhibir [categoría] sin sacrificar otras', 'Espacio', 'Post + carrusel'),
         ('111', '"Lo venden más barato en [competidor]": cómo responder', 'Competitividad', 'Post'),
         ('112', '"No conozco la marca": por qué [marca] es una apuesta segura', 'Marcas nuevas', 'Carrusel'),
         ('113', '"¿Y la garantía?": lo que necesitas saber para vender [marca] con confianza', 'Garantía', 'Post + carrusel'),
         ('114', '"Ya tengo proveedor": por qué tener dos siempre conviene', 'Segunda fuente', 'Post'),
         ('115', '"No sé cómo venderlo": la capacitación que te damos para [categoría]', 'Soporte', 'Carrusel'),
     ]),
    ('ARQUETIPO 12: SEGMENTACIÓN POR VERTICAL',
     'Fórmula: tipo de cliente final + necesidad + producto SUPRICOM',
     [
         ('116', 'Lo que un [oficina/empresa] necesita en [categoría] y tú le puedes vender', 'Segmentos', 'Carrusel'),
         ('117', '[Categoría] para [hoteles/clínicas/escuelas]: lo que están comprando', 'Verticales', 'Carrusel'),
         ('118', 'El kit de [categoría] que toda [pequeña empresa] debería tener', 'Segmento', 'Post + carrusel'),
         ('119', '¿Tienes clientes en [sector]? Esto es lo que más necesitan', 'Por industria', 'Carrusel'),
         ('120', '[Sector]: la vertical que más está invirtiendo en [categoría] este año', 'Oportunidad', 'Post + carrusel'),
         ('121', '5 tipos de clientes que necesitan [producto] y quizás no los ves', 'Descubrimiento', 'Carrusel'),
         ('122', 'Lo que las [empresas de X] buscan en un [producto]', 'Buyer persona', 'Carrusel 8 slides'),
         ('123', '[Vertical] está creciendo: ¿cómo capturar ese cliente con [categoría]?', 'Crecimiento', 'Post'),
         ('124', 'Por qué un [sector] prefiere [marca] sobre otras opciones', 'Lealtad', 'Post'),
         ('125', 'El cliente que no sabías que existía: [sector] necesita [producto]', 'Nuevos', 'Carrusel'),
     ]),
    ('ARQUETIPO 13: CASOS DE ÉXITO / "STORYTELLING"',
     'Fórmula: situación + solución + resultado',
     [
         ('126', 'Cómo una tienda duplicó sus ventas de [categoría] en un mes', 'Categorías', 'Carrusel'),
         ('127', 'El cliente que no quería [producto] y terminó comprando [X] unidades', 'Persuasión', 'Post + carrusel'),
         ('128', 'De vender [X] a [Y] en 30 días: lo que hizo este revendedor con [categoría]', 'Escalabilidad', 'Carrusel'),
         ('129', 'La tienda que apostó por [marca] y hoy es su mejor canal', 'Marcas', 'Post + carrusel'),
         ('130', 'Cómo una categoría "olvidada" se volvió la más rentable', 'Subestimadas', 'Carrusel'),
         ('131', 'El error que este revendedor cometió con [producto] (y cómo lo corrigió)', 'Lección', 'Post'),
         ('132', 'Lo que este dueño de tienda dice ahora sobre [categoría] (y que antes no sabía)', 'Testimonial', 'Reel 30s'),
         ('133', 'De 0 a [X] ventas semanales: el plan que usó esta tienda para [categoría]', 'Metodología', 'Carrusel 10 slides'),
         ('134', 'La vez que [producto] salvó a un negocio de perder [X]$', 'Problema/solución', 'Post'),
         ('135', '3 tiendas, 3 estrategias, 1 resultado: cómo cada una vende [categoría]', 'Diversidad', 'Carrusel'),
     ]),
    ('ARQUETIPO 14: BENEFICIO SIMPLE / "SIN TECNICISMOS"',
     'Fórmula: beneficio concreto + ausencia de specs + aplicabilidad',
     [
         ('136', '¿Para qué sirve [producto] en la vida real? (explicado simple)', 'Técnicos', 'Reel 20s'),
         ('137', 'Lo que [producto] hace por tu cliente sin que expliques specs', 'Cualquiera', 'Post'),
         ('138', 'En 10 segundos: por qué tu cliente necesita [producto]', 'Cualquiera', 'Reel 10s'),
         ('139', 'Lo que le pasa a [situación] si NO tiene [producto]', 'Protección', 'Reel 15s'),
         ('140', '[Producto] en 3 palabras: [b1], [b2], [b3]', 'Cualquiera', 'Post'),
         ('141', 'Una imagen que explica por qué [categoría] es importante hoy', 'Críticas', 'Imagen'),
         ('142', 'Cómo explicar [producto] a un cliente que no entiende de tecnología', 'Cualquiera', 'Carrusel'),
         ('143', 'Esto no es un gasto, es una inversión: [producto]', 'Valor', 'Post'),
         ('144', '[X] razones para agradecerte que recomiendes [producto]', 'Cualquiera', 'Carrusel'),
         ('145', 'Lo simple de [producto]: enchufa y funciona', 'Plug-and-play', 'Post + reel'),
     ]),
    ('ARQUETIPO 15: INTERACCIÓN / "ENGAGEMENT DIRECTO"',
     'Fórmula: pregunta + participación + comunidad',
     [
         ('146', '¿Ya vendes [categoría]? Cuéntanos cuál es el que más te piden', 'Cualquiera', 'Post encuesta'),
         ('147', '¿Cuál de estos [producto] crees que se vende más?', 'Cualquiera', 'Post votación'),
         ('148', '¿Qué objeción escuchas más al vender [producto]?', 'Interacción', 'Post'),
         ('149', '¿Cuánto crees que cuesta [producto]? Adivina', 'Gamificación', 'Post + Stories'),
         ('150', 'Cuéntanos: ¿cuál fue tu mayor venta de [categoría] esta semana?', 'Comunidad', 'Post'),
         ('151', '¿Prefieres vender [marca A] o [marca B]? ¿Por qué?', 'Marcas', 'Post'),
         ('152', 'Marca la opción: ¿precio o calidad al vender [producto]?', 'Encuesta', 'Post'),
         ('153', 'Recomiéndale [categoría] a otro dueño de tienda', 'Peer-to-peer', 'Post'),
         ('154', '¿Qué te gustaría aprender a vender mejor este mes?', 'Feedback', 'Post'),
         ('155', 'De dueño a dueño: ¿qué [producto] te ha dado mejor resultado?', 'Testimonial', 'Post'),
     ]),
    ('ARQUETIPO 16: CONTENIDO GENERATIVO / "IA Y FUTURO"',
     'Fórmula: tendencia tecnológica + implicación para la tienda',
     [
         ('156', 'Cómo la IA está cambiando lo que tus clientes piden en [categoría]', 'Cualquiera', 'Post + carrusel'),
         ('157', 'Lo que los compradores B2B buscan en ChatGPT (y cómo aparecer ahí)', 'SUPRICOM', 'Carrusel'),
         ('158', '¿Tus clientes investigan antes de comprar? Esto encuentran', 'Categorías', 'Carrusel'),
         ('159', 'El futuro de [categoría]: hacia dónde va el mercado', 'Proyección', 'Carrusel 10 slides'),
         ('160', 'Las [categoría] del futuro que ya están aquí', 'Innovación', 'Post + carrusel'),
     ]),
    ('ARQUETIPO 17: LIQUIDACIÓN Y OPORTUNIDAD',
     'Fórmula: oportunidad + tiempo limitado + acción',
     [
         ('161', 'Últimas unidades de [producto] — no sabemos cuándo vuelven', 'Stock bajo', 'Post + Stories'),
         ('162', '[Producto] a precio especial esta semana', 'Cualquiera', 'Post'),
         ('163', 'Lote especial de [categoría] con condiciones preferenciales', 'Categorías', 'Post + carrusel'),
         ('164', 'Oferta por volumen en [producto]: mientras más, mejor precio', 'Margen', 'Post'),
         ('165', 'Descuento por pronta compra en [categoría] — solo [X] días', 'Categorías', 'Post + Stories'),
     ]),
    ('ARQUETIPO 18: PREGUNTAS FRECUENTES',
     'Fórmula: duda común + respuesta clara + CTA',
     [
         ('166', 'Lo que todo dueño de tienda pregunta antes de comprar [categoría]', 'Categorías', 'Carrusel'),
         ('167', 'Las [X] preguntas más frecuentes sobre [producto] en SUPRICOM', 'Productos', 'Carrusel 8 slides'),
         ('168', 'Mitos y verdades sobre [categoría] que todo revendedor debe conocer', 'Con mitos', 'Carrusel'),
         ('169', '¿Cómo se compara [producto] con lo que ya tienes?', 'Sustitución', 'Post'),
         ('170', 'Lo que nadie te dice de [categoría] y nosotros sí', 'Transparencia', 'Carrusel'),
     ]),
    ('ARQUETIPO 19: INSPIRACIONAL / "EL POTENCIAL"',
     'Fórmula: visión + posibilidad + llamado a crecer',
     [
         ('171', 'Si solo pudieras vender una [categoría] esta semana, que sea esta', 'Priorización', 'Post'),
         ('172', 'El potencial de [categoría] que la mayoría de tiendas ignora', 'Oportunidad', 'Carrusel'),
         ('173', 'Lo que [marca] puede hacer por tu tienda', 'Visión partner', 'Post'),
         ('174', 'Así se ve una tienda que entiende [categoría]', 'Aspiracional', 'Carrusel'),
         ('175', '¿Y si este mes pruebas con [categoría]?', 'Reto', 'Post'),
     ]),
    ('ARQUETIPO 20: RECORDATORIO / "NO OLVIDES"',
     'Fórmula: necesidad básica + producto esencial',
     [
         ('176', 'No olvides tener [producto] en tu inventario esta semana', 'Básicos', 'Post'),
         ('177', 'Checklist semanal de [categoría]: lo que no puede faltar', 'Lista', 'Post imagen'),
         ('178', 'Revisa tu stock de [categoría]: lo que deberías tener según ventas', 'Planificación', 'Post + carrusel'),
         ('179', 'Si tienes clientes de [sector], ofréceles [producto]', 'Recordatorio', 'Post'),
         ('180', '3 [producto] que siempre deberías tener en tu vitrina', 'Imprescindibles', 'Post'),
     ]),
]

for titulo, formula, temas in arquetipos:
    doc.add_page_break()
    add_heading_styled(doc, titulo, 2)
    p = doc.add_paragraph()
    run = p.add_run(formula)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    rows = [['#', 'Título genérico', 'Aplica a', 'Formato']]
    for num, tema, aplica, fmt in temas:
        rows.append([num, tema, aplica, fmt])
    add_table_with_shading(doc, rows)

# ════════════════════════════════════════════
# MATRIZ DE APLICACIÓN RÁPIDA
# ════════════════════════════════════════════
doc.add_page_break()
add_heading_styled(doc, 'MATRIZ DE APLICACIÓN RÁPIDA', 1)
p = doc.add_paragraph()
p.add_run('Para usar esta base de datos, toma cualquier arquetipo + cualquier variable:')

rows = [['Variable', 'Ejemplos']]
for var, ej in [
    ('[producto]', 'UPS, regulador, impresora, router, laptop, power bank, BLUETTI, SSD'),
    ('[categoría]', 'protección eléctrica, almacenamiento, redes, impresión, audio, tablets'),
    ('[marca]', 'FORZA, APC, SMARTBITT, BLUETTI, HP, HIKVISION, KINGSTON, ADATA'),
    ('[mes]', 'julio, agosto, septiembre'),
    ('[situación país]', 'crisis eléctrica, temporada de lluvias, apertura económica'),
    ('[sector/vertical]', 'oficinas, clínicas, hoteles, escuelas, talleres, comercios'),
    ('[competidor]', 'KODE, Magnabyte, SARVEN, PC Suplidores, GAHL Tech'),
    ('[%]', 'dato de crecimiento o margen'),
    ('[X] unidades', 'cifra de stock o ventas'),
]:
    rows.append([var, ej])
add_table_with_shading(doc, rows)

# ════════════════════════════════════════════
# FRECUENCIA SUGERIDA
# ════════════════════════════════════════════
add_heading_styled(doc, 'FRECUENCIA SUGERIDA POR ARQUETIPO (semanal)', 1)
rows = [['Arquetipo', 'Posts/sem', 'Propósito']]
for a, psts, prop in [
    ('Prueba social / Lo que se vende', '2', 'Urgencia + validación'),
    ('Caso de negocio / Por qué tenerlo', '1', 'Conversión'),
    ('Educación / Cómo venderlo', '2', 'Autoridad + utilidad'),
    ('Inteligencia de mercado', '1', 'Pensamiento líder'),
    ('Stock y disponibilidad', '1', 'Operacional'),
    ('Comparativas', '1', 'Decisión'),
    ('Estacional', '1', 'Timing'),
    ('Detrás del mayorista', '1', 'Confianza'),
    ('Rentabilidad', '1', 'ROI'),
    ('Kits y combos', '1', 'Ticket promedio'),
    ('Objeciones', '1', 'Cierre'),
    ('Interacción', '1', 'Comunidad'),
]:
    rows.append([a, psts, prop])
add_table_with_shading(doc, rows)

p = doc.add_paragraph()
run = p.add_run('Total semanal aprox: 14 posts (2/día promedio)')
run.bold = True

# ════════════════════════════════════════════
# PLANTILLA DE ADAPTACIÓN RÁPIDA
# ════════════════════════════════════════════
doc.add_page_break()
add_heading_styled(doc, 'PLANTILLA DE ADAPTACIÓN RÁPIDA', 1)

p = doc.add_paragraph()
run = p.add_run('Para convertir cualquier tema genérico en contenido concreto:')
run.bold = True

doc.add_paragraph('Paso 1: Elegir arquetipo + número de tema', style='List Number')
doc.add_paragraph('Paso 2: Reemplazar variables entre [brackets]', style='List Number')
doc.add_paragraph('Paso 3: Elegir formato (post, carrusel, reel)', style='List Number')
doc.add_paragraph('Paso 4: Aplicar framework correspondiente:', style='List Number')

frameworks = [
    'HAVPC para video (Hook, Agitation, Visualization, Proof, CTA)',
    'Swipe to Learn para carrusel (título → problema → solución → datos → CTA)',
    'AGES para distribución (30% Awareness, 30% General Interest, 25% Education, 15% Sales)',
]
for fw in frameworks:
    doc.add_paragraph(fw, style='List Bullet')

doc.add_paragraph()

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Documento generado Julio 2026 —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Base: 180 temas genéricos × 23 categorías = 4,140 combinaciones posibles')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SUPRICOM. Tu mayorista de confianza.')
run.font.color.rgb = RGBColor(0xCC, 0x7A, 0x00)
run.bold = True

# ── Save ──
output = r'C:\Users\MDIGITAL01\Desktop\MARKETING OPENCODE\base_temas_genericos_SUPRICOM.docx'
doc.save(output)
print(f'Documento guardado en: {output}')
