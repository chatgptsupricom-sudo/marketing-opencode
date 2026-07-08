from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_table_with_shading(doc, rows_data, col_widths=None):
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                paragraph.space_after = Pt(2)
                paragraph.space_before = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A1A2E"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()
    return table

def add_bullet(doc, text, bold_prefix=None, indent=None):
    p = doc.add_paragraph(style='List Bullet')
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_number(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('MI GRAN IDEA')
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('PROGRAMA DE INNOVACIÓN INTERNA')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xCC, 0x7A, 0x00)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('Buzón Digital de Sugerencias Integrado al Panel SUPRICOM')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = line.add_run('—' * 40)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Propuesta preparada para Gerencia de Ventas')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('Julio 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SUPRICOM  ·  Tu mayorista de confianza')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xCC, 0x7A, 0x00)
run.bold = True

doc.add_page_break()

# ════════════════════════════════════════════
# TABLA DE CONTENIDO
# ════════════════════════════════════════════
add_heading_styled(doc, 'CONTENIDO', 1)
toc_items = [
    'Resumen Ejecutivo',
    '¿Qué es "Mi Gran Idea Supricom"?',
    'Por qué Digital y no Físico',
    'Integración con el Panel SUPRICOM',
    'Funcionalidades del Módulo',
    'Flujo del Proceso (Paso a Paso)',
    'Categorías de Ideas',
    'Sistema de Evaluación y Puntaje',
    'Sistema de Reconocimiento',
    'Roles y Responsabilidades',
    'KPIs del Programa',
    'Plan de Implementación',
    'Anexos',
]
for i, item in enumerate(toc_items, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {item}')
    run.font.size = Pt(11)
    if i < 10:
        p.paragraph_format.left_indent = Cm(0.5)

doc.add_page_break()

# ════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ════════════════════════════════════════════
add_heading_styled(doc, '1. Resumen Ejecutivo', 1)

doc.add_paragraph(
    '"Mi Gran Idea Supricom" es un programa permanente de innovación interna '
    'que canaliza la creatividad de todos los colaboradores de SUPRICOM para '
    'proponer mejoras en procesos, productos, servicio al cliente y ambiente laboral.'
)
doc.add_paragraph(
    'A diferencia de los buzones de sugerencias tradicionales, esta propuesta '
    'se integra directamente al Panel SUPRICOM (actualmente en desarrollo por '
    'el equipo de programación), convirtiéndose en un módulo nativo al que '
    'todo el personal accede con sus credenciales existentes.'
)

p = doc.add_paragraph()
run = p.add_run('Principios del programa:')
run.bold = True

add_bullet(doc, ' Sin presupuesto inicial — aprovecha infraestructura existente')
add_bullet(doc, ' Cobertura total — 100% del personal (oficina, campo, sucursales)')
add_bullet(doc, ' Trazabilidad completa — cada idea tiene un ciclo de vida visible')
add_bullet(doc, ' Anonimato configurable — quien quiera participar sin identificarse, puede')
add_bullet(doc, ' Sin fricción — no más papeles, ni urnas, ni digitalización manual')

doc.add_page_break()

# ════════════════════════════════════════════
# 2. QUÉ ES "MI GRAN IDEA SUPRICOM"
# ════════════════════════════════════════════
add_heading_styled(doc, '2. ¿Qué es "Mi Gran Idea Supricom"?', 1)

doc.add_paragraph(
    'Es el buzón de sugerencias oficial de SUPRICOM, pero en su versión moderna, '
    'digital y conectada. No se trata de una urna física donde las ideas desaparecen, '
    'sino de un sistema donde cada propuesta tiene un ciclo de vida:'
)
add_bullet(doc, ' Registro de la idea con categoría y descripción', '1. ')
add_bullet(doc, ' Evaluación con criterios objetivos', '2. ')
add_bullet(doc, ' Decisión: aprobada, en estudio o devuelta con motivo', '3. ')
add_bullet(doc, ' Implementación con seguimiento', '4. ')
add_bullet(doc, ' Reconocimiento al autor', '5. ')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Objetivos del programa:')
run.bold = True

add_number(doc, ' Fomentar una cultura de mejora continua e innovación participativa')
add_number(doc, ' Identificar oportunidades de ahorro, eficiencia y nuevos ingresos que la gerencia no ve desde su posición')
add_number(doc, ' Dar voz a todos los colaboradores, independientemente de su rol o ubicación')
add_number(doc, ' Reconocer y premiar el talento interno que genera valor medible')

doc.add_page_break()

# ════════════════════════════════════════════
# 3. POR QUÉ DIGITAL Y NO FÍSICO
# ════════════════════════════════════════════
add_heading_styled(doc, '3. Por qué Digital y no Físico', 1)

doc.add_paragraph(
    'Se evaluó la opción de buzones físicos (urnas en oficina) y se descartó '
    'por las siguientes razones fundamentadas en la literatura de engagement '
    'organizacional y en la realidad operativa de SUPRICOM:'
)

rows = [
    ['Factor', 'Buzón Físico', 'Buzón Digital (Panel SUPRICOM)'],
    ['Cobertura', 'Solo personal presente en oficina', '100% del personal (oficina, campo, sucursales, remoto)'],
    ['Costo operativo', 'Impresión de boletas, bolígrafos, urna + 2-3 hrs/semana para abrir, leer, descifrar y digitalizar', '$0. Automatizado. El sistema entrega datos listos para evaluar.'],
    ['Anonimato real', 'No. La caligrafía, ubicación y hora delatan al autor.', 'Sí. Configurable sin registro de IP ni datos personales.'],
    ['Trazabilidad', 'Nula. La idea entra al buzón y no se sabe qué pasó. El empleado nunca recibe respuesta.', 'Total. Cada idea tiene estatus visible: recibida, en evaluación, aprobada, implementada, rechazada (con motivo).'],
    ['Categorización', 'Manual. Alguien debe leer y clasificar cada papel.', 'Automática. El empleado selecciona categoría de una lista desplegable. Datos limpios desde el origen.'],
    ['Escalabilidad', 'Mala. Cada sucursal necesitaría su propio buzón.', 'Inmediata. El Panel es único y accesible desde cualquier ubicación.'],
    ['Fomento de participación', 'Bajo. El "agujero negro" desmotiva.', 'Alto. El empleado ve que su idea fue recibida, evaluada y recibe respuesta.'],
    ['Análisis de datos', 'Imposible. Datos en papel no agregables.', 'Inmediato. Reportes por categoría, departamento, mes, estatus.'],
]
add_table_with_shading(doc, rows)

p = doc.add_paragraph()
run = p.add_run('\nConclusión: ')
run.bold = True
p.add_run(
    'El buzón digital es superior en todos los factores críticos. '
    'Implementar un buzón físico sería más costoso de lo que parece a simple vista '
    '(sin presupuesto asignado), excluiría a una parte importante del equipo '
    'y generaría frustración por falta de retroalimentación.'
)

doc.add_page_break()

# ════════════════════════════════════════════
# 4. INTEGRACIÓN CON EL PANEL SUPRICOM
# ════════════════════════════════════════════
add_heading_styled(doc, '4. Integración con el Panel SUPRICOM', 1)

doc.add_paragraph(
    'El Panel SUPRICOM es la plataforma interna que el equipo de programación '
    'se encuentra desarrollando actualmente. Es el sitio donde todos los colaboradores '
    'deberán ingresar para acceder a herramientas, información y recursos de la empresa.'
)

doc.add_paragraph(
    '"Mi Gran Idea Supricom" se integrará como un módulo nativo dentro del Panel, '
    'con las siguientes ventajas frente a una solución externa (Google Forms, Jotform, etc.):'
)

add_number(doc, ' Autenticación unificada — el colaborador ingresa con sus credenciales del Panel. No requiere registro adicional.')
add_number(doc, ' Experiencia integrada — no hay que salir del Panel, abrir otra página ni recordar otro enlace.')
add_number(doc, ' Datos centralizados — las ideas se almacenan en la misma base de datos del Panel.')
add_number(doc, ' Notificaciones internas — el Panel puede mostrar alertas de nuevas ideas, cambios de estatus, resultados mensuales.')
add_number(doc, ' Perfil de participante — el Panel registra el historial de ideas de cada colaborador (ideas enviadas, aprobadas, implementadas).')

doc.add_page_break()

# ════════════════════════════════════════════
# 5. FUNCIONALIDADES DEL MÓDULO
# ════════════════════════════════════════════
add_heading_styled(doc, '5. Funcionalidades del Módulo', 1)

add_heading_styled(doc, '5.1 Para el colaborador (usuario)', 2)

add_bullet(doc, 'Botón visible "Enviar mi idea" en el menú principal del Panel.')
add_bullet(doc, 'Formulario simple con los siguientes campos:')
add_bullet(doc, ' Departamento (Ventas, Marketing, Operaciones, Logística, Administración, Sistemas, Otro)', indent=1)
add_bullet(doc, ' Categoría de la idea (lista desplegable — ver sección 7)', indent=1)
add_bullet(doc, ' Título corto de la idea (máx. 100 caracteres)', indent=1)
add_bullet(doc, ' Descripción: ¿Cuál es el problema u oportunidad que detectaste? (campo abierto)', indent=1)
add_bullet(doc, ' Propuesta de solución: ¿Qué propones hacer? (campo abierto)', indent=1)
add_bullet(doc, ' Beneficio esperado: ¿Cómo ayuda esto a SUPRICOM? (selección: Ahorro de costos / Incremento de ingresos / Mejora de proceso / Mejora de clima laboral / Otro)', indent=1)
add_bullet(doc, ' ¿Participarías en la implementación? (Sí / No)', indent=1)

add_bullet(doc, 'Opción de anonimato:')
add_bullet(doc, ' Casilla: "Enviar de forma anónima (no optarás al premio)"', indent=1)
add_bullet(doc, ' Si elige anónimo, el sistema no asocia la idea a su usuario.', indent=1)

add_bullet(doc, 'Historial personal:')
add_bullet(doc, ' Sección "Mis Ideas" donde el colaborador ve el estatus de cada una de sus propuestas.', indent=1)
add_bullet(doc, ' Estatus visibles: Recibida → En Evaluación → Aprobada → Implementada → Rechazada (con motivo)', indent=1)

add_bullet(doc, 'Tablero público de innovación:')
add_bullet(doc, ' Sección visible para todos con las ideas implementadas del mes.', indent=1)
add_bullet(doc, ' Ranking de participación por departamento (sin nombres si es anónimo).', indent=1)
add_bullet(doc, ' Foto del "Innovador del Mes" (si acepta visibilidad).', indent=1)

add_heading_styled(doc, '5.2 Para el administrador / comité evaluador', 2)

add_bullet(doc, 'Panel de administración con:')
add_bullet(doc, ' Bandeja de entrada de nuevas ideas (ordenadas por fecha).', indent=1)
add_bullet(doc, ' Filtros por: categoría, departamento, fecha, estatus.', indent=1)
add_bullet(doc, ' Visualización detallada de cada idea.', indent=1)

add_bullet(doc, 'Herramientas de evaluación:')
add_bullet(doc, ' Asignación de puntaje según matriz de 100 puntos (ver sección 8).', indent=1)
add_bullet(doc, ' Campo de comentarios internos del comité.', indent=1)
add_bullet(doc, ' Botón para cambiar estatus: Aprobada / En Estudio / Rechazada (campo obligatorio de motivo).', indent=1)

add_bullet(doc, 'Reportes automáticos:')
add_bullet(doc, ' Ideas recibidas por mes (gráfico de barras).', indent=1)
add_bullet(doc, ' % de ideas implementadas.', indent=1)
add_bullet(doc, ' Participación por departamento.', indent=1)
add_bullet(doc, ' Ideas por categoría.', indent=1)
add_bullet(doc, ' Tiempo promedio de evaluación.', indent=1)

add_bullet(doc, 'Exportación a Excel de todo el histórico para análisis profundos.')

doc.add_page_break()

# ════════════════════════════════════════════
# 6. FLUJO DEL PROCESO
# ════════════════════════════════════════════
add_heading_styled(doc, '6. Flujo del Proceso (Paso a Paso)', 1)

rows = [
    ['Paso', 'Qué ocurre', 'Responsable', 'Tiempo'],
    ['1', 'El colaborador ingresa al Panel, hace clic en "Enviar mi idea" y completa el formulario.', 'Colaborador', '5-10 min'],
    ['2', 'El sistema registra la idea, asigna un número de folio automático y envía notificación al comité evaluador.', 'Panel (automático)', 'Inmediato'],
    ['3', 'El colaborador recibe confirmación: "Tu idea #123 fue recibida. Gracias por contribuir."', 'Panel (automático)', 'Inmediato'],
    ['4', 'El comité evaluador revisa las ideas acumuladas. Evalúa cada una con la matriz de puntaje.', 'Comité evaluador', 'Días 26-28 de cada mes'],
    ['5', 'El sistema actualiza el estatus de cada idea: Aprobada, En Estudio o Rechazada (con motivo visible para el autor).', 'Comité evaluador', 'Día 28'],
    ['6', 'Las ideas aprobadas pasan a fase de implementación. Se asigna responsable y fecha estimada.', 'Gerencia / Jefe de área', 'Según complejidad'],
    ['7', 'El último día del mes: publicación de resultados, idea del mes, premios y actualización del tablero público.', 'Comité evaluador + Marketing', 'Día 30/31'],
    ['8', 'Seguimiento trimestral: las ideas implementadas se evalúan por su impacto real (ahorro, ingresos, mejora medible).', 'Gerencia General', 'Cada 3 meses'],
]
add_table_with_shading(doc, rows)

p = doc.add_paragraph()
run = p.add_run('Nota sobre el ciclo mensual: ')
run.bold = True
p.add_run(
    'Las ideas se recolectan de forma continua (días 1-25), se evalúan '
    'en bloque al cierre (días 26-28) y se comunican los resultados el '
    'último día hábil del mes. Esto evita la fatiga de evaluar una por una '
    'en tiempo real y genera un "evento mensual" que mantiene el interés.'
)

doc.add_page_break()

# ════════════════════════════════════════════
# 7. CATEGORÍAS DE IDEAS
# ════════════════════════════════════════════
add_heading_styled(doc, '7. Categorías de Ideas', 1)

doc.add_paragraph(
    'Cada idea debe pertenecer a una de las siguientes categorías. '
    'Esto permite filtrar, priorizar y medir el impacto por área de negocio.'
)

rows = [
    ['Categoría', 'Descripción', 'Ejemplos de ideas'],
    ['Ahorro de Costos', 'Propuestas que reducen gastos operativos, logísticos o administrativos.', 'Reducción de consumo eléctrico en almacén, renegociación de fletes, optimización de empaques.'],
    ['Mejora de Procesos', 'Ideas que optimizan flujos de trabajo, reducen tiempos o eliminan pasos innecesarios.', 'Automatización de órdenes de compra, mejora en el proceso de picking, digitalización de documentos.'],
    ['Nuevos Productos / Mercados', 'Propuestas para expandir el catálogo, ingresar a nuevas categorías o mercados.', 'Traer una nueva marca de impresoras, abrir canal de venta a Panamá, lanzar línea de accesorios gaming.'],
    ['Servicio al Cliente', 'Ideas para mejorar la experiencia de los aliados (tiendas) en su relación con SUPRICOM.', 'Chatbot en WhatsApp, notificaciones de tracking de envío, programa de fidelización.'],
    ['Ambiente Laboral', 'Propuestas para mejorar el clima organizacional, la comunicación interna o el bienestar del equipo.', 'Día de integración trimestral, implementar horario flexible, mejorar el área de descanso.'],
    ['Tecnología / Sistemas', 'Mejoras al Panel SUPRICOM, CRM, sistemas internos o herramientas digitales.', 'Nuevo reporte en el CRM, integración con WhatsApp Business API, mejora en la interfaz del Panel.'],
    ['Marketing / Ventas', 'Ideas para mejorar la comunicación externa, la marca o las estrategias de venta.', 'Campaña estacional para Q4, nuevo formato de catálogo digital, estrategia de email marketing a aliados.'],
    ['Otro', 'Ideas que no encajan en las categorías anteriores.', 'Cualquier propuesta no clasificable en las categorías previas.'],
]
add_table_with_shading(doc, rows)

doc.add_page_break()

# ════════════════════════════════════════════
# 8. SISTEMA DE EVALUACIÓN
# ════════════════════════════════════════════
add_heading_styled(doc, '8. Sistema de Evaluación y Puntaje', 1)

doc.add_paragraph(
    'Cada idea es evaluada por el comité usando la siguiente matriz de 100 puntos. '
    'Este sistema garantiza objetividad y criterios uniformes para todas las propuestas.'
)

rows = [
    ['Criterio', 'Puntaje Máximo', '¿Qué evalúa?'],
    ['Impacto para SUPRICOM', '30 pts', 'Ahorro económico, incremento de ingresos o mejora del servicio al cliente que genera la idea.'],
    ['Facilidad de Implementación', '20 pts', 'Recursos necesarios, tiempo de ejecución y complejidad técnica para llevar la idea a la realidad.'],
    ['Innovación y Creatividad', '15 pts', 'Originalidad de la propuesta y qué tan diferente es a lo que ya se hace o se ha intentado.'],
    ['Alcance', '10 pts', '¿La idea beneficia a una persona, un departamento o a toda la empresa?'],
    ['Sostenibilidad', '10 pts', '¿Los beneficios son permanentes y escalables, o solo un evento puntual?'],
    ['Riesgo', '5 pts', 'Riesgo operativo, financiero o legal de implementar la idea. A menor riesgo, mayor puntaje.'],
    ['Alineación Estratégica', '10 pts', 'Coherencia con la visión, misión y objetivos estratégicos de SUPRICOM.'],
    ['TOTAL', '100 pts', '—'],
]
add_table_with_shading(doc, rows)

p = doc.add_paragraph()
run = p.add_run('Puntaje mínimo para aprobación: ')
run.bold = True
p.add_run('60 puntos. Por debajo de 60, la idea pasa a "En Estudio" o se rechaza con motivo.')

doc.add_page_break()

# ════════════════════════════════════════════
# 9. SISTEMA DE RECONOCIMIENTO
# ════════════════════════════════════════════
add_heading_styled(doc, '9. Sistema de Reconocimiento', 1)

doc.add_paragraph(
    'Para mantener el programa vivo, el reconocimiento debe ser predecible, '
    'transparente y escalonado. Se proponen tres niveles:'
)

add_heading_styled(doc, 'Nivel 1: Participación Mensual (todos los que envían una idea válida)', 2)
add_bullet(doc, ' Sorteo de un detalle corporativo (termo, gorra, cupón de café) entre todos los participantes del mes.')
add_bullet(doc, ' Esto premia el intento y mantiene el hábito, incluso si la idea no fue seleccionada.')

add_heading_styled(doc, 'Nivel 2: La Gran Idea del Mes (1 ganador mensual)', 2)
add_bullet(doc, ' La idea mejor puntuada por el comité recibe un incentivo intermedio.')
add_bullet(doc, ' Opciones de premio (a confirmar según disponibilidad): bono en efectivo, tarde libre remunerada, tarjeta de regalo, o producto del catálogo SUPRICOM.', indent=1)
add_bullet(doc, ' Publicación en el tablero de innovación y comunicación interna.')

add_heading_styled(doc, 'Nivel 3: El Innovador del Año (1 ganador anual)', 2)
add_bullet(doc, ' Entre las 12 ideas del mes (o las implementadas durante el año), se elige la de mayor impacto económico o cultural para SUPRICOM.')
add_bullet(doc, ' Premio mayor: viaje corto nacional, electrodoméstico, bono económico significativo, o producto de alto valor del catálogo.')
add_bullet(doc, ' Reconocimiento público en la comunicación interna y mención especial del Gerente General.')

doc.add_page_break()

# ════════════════════════════════════════════
# 10. ROLES Y RESPONSABILIDADES
# ════════════════════════════════════════════
add_heading_styled(doc, '10. Roles y Responsabilidades', 1)

rows = [
    ['Rol', 'Responsabilidades', '¿Quién lo ocupa?'],
    ['Sponsor', 'Aprueba el programa. Anuncia el lanzamiento. Comunica resultados a la organización. Valida ideas del área comercial.', 'Gerente de Ventas'],
    ['Dueño Operativo', 'Diseña el formulario en el Panel. Redacta las comunicaciones internas. Gestiona el tablero público. Da soporte técnico a los usuarios.', 'Marketing / Sistemas'],
    ['Comité Evaluador', 'Revisa ideas mensualmente. Asigna puntaje según la matriz. Determina aprobación, rechazo o pase a estudio. Selecciona la Idea del Mes.', 'Gerente de Ventas + Marketing + RRHH + 1 miembro rotativo de Operaciones'],
    ['Implementador', 'Ejecuta la idea aprobada. Reporta avances. Documenta el resultado.', 'Jefe del área correspondiente'],
    ['Comunicador', 'Publica resultados mensuales. Actualiza el tablero. Mantiene el programa visible y presente.', 'Marketing'],
]
add_table_with_shading(doc, rows)

add_heading_styled(doc, '10.1 Modelo de Gobierno: ¿Quién manda en esto?', 2)

doc.add_paragraph(
    'El programa necesita un modelo de gobierno claro para no morir en el mes 2. '
    'Cada rol tiene responsabilidades distintas y complementarias:'
)

add_heading_styled(doc, 'Marketing — Dueño Operativo (el que ejecuta)', 3)
doc.add_paragraph(
    'Marketing es el departamento que debe liderar la operación del programa por las siguientes razones:'
)
add_bullet(doc, ' Ya maneja la comunicación interna y externa de SUPRICOM — nadie más tiene la vocería para anunciar, recordar y celebrar.')
add_bullet(doc, ' Sabe diseñar formularios, flujos de comunicación y piezas visuales — RRHH y Ventas no tienen esas competencias.')
add_bullet(doc, ' Entiende de segmentación — puede categorizar ideas y darles el tratamiento adecuado.')
add_bullet(doc, ' Tiene la relación con Sistemas para coordinar el desarrollo del módulo en el Panel.')
add_bullet(doc, ' Puede medir y reportar — los KPIs del programa son análogos a los KPIs de cualquier campaña de marketing interno.')

add_heading_styled(doc, 'Gerente de Ventas — Sponsor (el que da la cara)', 3)
doc.add_paragraph(
    'El Gerente de Ventas es el solicitante original del programa y debe ser su patrocinador visible. '
    'Su rol es:'
)
add_bullet(doc, ' Anunciar el lanzamiento del programa con su autoridad — da peso y seriedad a la iniciativa.')
add_bullet(doc, ' Validar ideas del área comercial — conoce qué problemas operativos son prioritarios.')
add_bullet(doc, ' Asistir a la premiación mensual — su presencia eleva el reconocimiento.')
add_bullet(doc, ' No gestionar el día a día — ese es trabajo de Marketing.')

add_heading_styled(doc, 'RRHH — Aliado Estratégico (el que evalúa personas)', 3)
doc.add_paragraph(
    'RRHH no debe operar el programa, pero sí participar en la evaluación porque:'
)
add_bullet(doc, ' Puede detectar sesgos en la participación (¿solo un departamento está participando?).')
add_bullet(doc, ' Asegura que el sistema de premios sea justo y transparente.')
add_bullet(doc, ' Gestiona el aspecto cultural — el programa es, al final, una herramienta de clima laboral.')

add_heading_styled(doc, 'Sistemas / Programación — Facilitador Técnico', 3)
doc.add_paragraph(
    'El equipo de programación tiene una responsabilidad acotada pero crítica:'
)
add_bullet(doc, ' Desarrollar el módulo dentro del Panel (formulario, base de datos, notificaciones, reportes).')
add_bullet(doc, ' Dar soporte técnico si hay fallas.')
add_bullet(doc, ' No participa en la evaluación ni en la comunicación del programa.')

add_heading_styled(doc, 'Estructura de Reporte', 3)

rows2 = [
    ['Rol', 'Reporta a', 'Lo nombra'],
    ['Sponsor (Gerente de Ventas)', 'Gerencia General', 'Su cargo'],
    ['Dueño Operativo (Marketing)', 'Sponsor', 'El Sponsor lo asigna'],
    ['Comité Evaluador', 'Sponsor + Gerencia General', 'El Sponsor lo convoca'],
    ['Implementador', 'Su jefe de área', 'Su cargo'],
]
add_table_with_shading(doc, rows2)

doc.add_paragraph(
    'Esta estructura asegura que nadie se pregunte "¿y esto de quién es?" — '
    'Marketing opera, Ventas patrocina, RRHH evalúa, Sistemas construye. '
    'Punto.'
)

add_heading_styled(doc, 'Comité evaluador: composición sugerida', 2)
doc.add_paragraph(
    'Se recomienda un comité de 4 personas, con reunión mensual de máximo 60 minutos:'
)
add_bullet(doc, ' 1 representante de Marketing (dueño operativo — preside la reunión)', bold_prefix='')
add_bullet(doc, ' 1 representante de RRHH (gestión cultural y transparencia)', bold_prefix='')
add_bullet(doc, ' 1 representante de Ventas (visión comercial y criterio de negocio)', bold_prefix='')
add_bullet(doc, ' 1 miembro rotativo de Operaciones o Logística (visión operativa, cambia cada 3 meses)', bold_prefix='')
doc.add_paragraph(
    'La Gerencia General participa solo en la aprobación de ideas que requieran '
    'presupuesto o cambio estructural significativo. No en la evaluación mensual.'
)

doc.add_page_break()

# ════════════════════════════════════════════
# 11. KPIs
# ════════════════════════════════════════════
add_heading_styled(doc, '11. KPIs del Programa', 1)

rows = [
    ['Indicador', 'Fórmula / Definición', 'Frecuencia', 'Meta sugerida (año 1)'],
    ['Ideas recibidas por mes', 'Conteo de ideas enviadas válidas', 'Mensual', '>= 10 ideas/mes'],
    ['% de participación', 'Colaboradores que envían al menos 1 idea / Total de colaboradores', 'Mensual', '>= 30%'],
    ['% de ideas implementadas', 'Ideas implementadas / Ideas aprobadas', 'Trimestral', '>= 60%'],
    ['Tiempo promedio de evaluación', 'Suma de días entre recepción y respuesta / Total de ideas', 'Mensual', '<= 30 días'],
    ['Ahorro económico generado', 'Suma de ahorros documentados de ideas implementadas', 'Trimestral', 'A definir según primeras ideas'],
    ['Participación por departamento', '% de colaboradores que participan por depto.', 'Mensual', '>= 20% en cada depto.'],
    ['Satisfacción del programa', 'Encuesta trimestral: "¿Sientes que tu idea fue escuchada?"', 'Trimestral', '>= 70% satisfacción'],
]
add_table_with_shading(doc, rows)

doc.add_page_break()

# ════════════════════════════════════════════
# 12. PLAN DE IMPLEMENTACIÓN
# ════════════════════════════════════════════
add_heading_styled(doc, '12. Plan de Implementación', 1)

rows = [
    ['Fase', 'Actividades', 'Responsable', 'Tiempo'],
    ['Fase 1: Diseño', 'Definir campos del formulario. Diseñar flujo de aprobación en el Panel. Definir matriz de evaluación.', 'Marketing + Sistemas', 'Semana 1'],
    ['Fase 2: Desarrollo', 'Programar módulo en el Panel. Crear base de datos de ideas. Configurar notificaciones automáticas.', 'Sistemas / Programadores', 'Semana 2-3'],
    ['Fase 3: Pruebas', 'Prueba interna del flujo completo. Ajustes de interfaz y funcionalidad.', 'Marketing + Sistemas', 'Semana 4'],
    ['Fase 4: Lanzamiento', 'Comunicación interna del programa. Correo del Gerente de Ventas anunciando "Mi Gran Idea". Publicación en el Panel.', 'Marketing + Gerente de Ventas', 'Semana 5'],
    ['Fase 5: Operación', 'Recolección continua. Primera evaluación al cierre del primer mes. Publicación de primeros resultados.', 'Comité evaluador', 'Mes 2 en adelante'],
]
add_table_with_shading(doc, rows)

doc.add_page_break()

# ════════════════════════════════════════════
# 13. ANEXOS
# ════════════════════════════════════════════
add_heading_styled(doc, '13. Anexos', 1)

add_heading_styled(doc, 'Anexo A: Referencias y mejores prácticas', 2)

doc.add_paragraph(
    'Este programa se basa en las siguientes referencias de engagement organizacional '
    'y gestión de innovación interna:'
)
add_bullet(doc, ' Qmarkets (2026): "The Employee Suggestion Box of the Future" — los buzones físicos tradicionales generan frustración cuando carecen de trazabilidad.')
add_bullet(doc, ' FaceUp (2025): "Ditch the Suggestion Box" — las plataformas digitales modernas ofrecen categorización, anonimato real y reporting que los buzones físicos no pueden igualar.')
add_bullet(doc, ' Sesame HR (2026): "Cómo implementar un buzón de sugerencias" — la tendencia global es migrar a digital por cobertura, costo y calidad de datos.')
add_bullet(doc, ' Teamflect (2025): "Employee Suggestion Box: 5 Best Free Apps" — las plataformas digitales permiten soporte multimedia, análisis de datos y trazabilidad.')

add_heading_styled(doc, 'Anexo B: Costo estimado del programa', 2)

rows = [
    ['Concepto', 'Costo'],
    ['Desarrollo del módulo en el Panel', '$0 (tiempo de programadores internos ya asignados al Panel)'],
    ['Hosting y almacenamiento', '$0 (infraestructura existente del Panel)'],
    ['Premios mensuales (12 meses)', 'A definir por Gerencia — puede ser $0 si se usan productos del catálogo SUPRICOM'],
    ['Premio anual', 'A definir por Gerencia'],
    ['Comunicación interna', '$0 (diseño in-house por Marketing)'],
    ['TOTAL estimado', 'SOLO premios (si aplica) — operación $0'],
]
add_table_with_shading(doc, rows)

add_heading_styled(doc, 'Anexo C: Ejemplo de comunicación de lanzamiento', 2)

doc.add_paragraph(
    'Asunto: "Lanzamos Mi Gran Idea Supricom — Tu voz crea el cambio"'
)
doc.add_paragraph(
    'Texto: "En SUPRICOM creemos que las mejores ideas nacen de quienes viven el día a día. '
    'Por eso lanzamos Mi Gran Idea Supricom, nuestro buzón digital de sugerencias integrado '
    'al Panel. Desde hoy puedes enviar tus propuestas para mejorar procesos, ahorrar costos, '
    'traer nuevos productos o hacer de SUPRICOM un mejor lugar para trabajar. '
    'Cada mes elegiremos la Gran Idea del Mes. Y a fin de año, el Innovador del Año. '
    'Ingresa al Panel, haz clic en "Enviar mi idea" y participa. Te leemos."'
)

doc.add_paragraph()

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
run = p.add_run('— Documento generado Julio 2026 —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SUPRICOM  ·  Programa Mi Gran Idea  ·  Versión 1.0')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SUPRICOM. Tu mayorista de confianza.')
run.font.color.rgb = RGBColor(0xCC, 0x7A, 0x00)
run.bold = True

# ── Save ──
output = r'C:\Users\MDIGITAL01\Desktop\MARKETING OPENCODE\Propuesta_Mi_Gran_Idea_SUPRICOM.docx'
doc.save(output)
print(f'Documento guardado en: {output}')
