from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)

PRIMARY = RGBColor(0x00, 0x56, 0xDB)
PRIMARY_LIGHT = RGBColor(0x07, 0x91, 0xF4)
TEXT = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x66, 0x66, 0x66)
DARK_BG = RGBColor(0x1A, 0x1A, 0x2E)

def add_slide_title(doc, num, title, fmt, pilar, products, ve_context):
    title_h = doc.add_heading(f'PIEZA #{num}: {title}', level=1)
    for run in title_h.runs:
        run.font.color.rgb = PRIMARY
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f'Formato: {fmt}  |  Pilar: {pilar}  |  Contexto VE: {ve_context}')
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = GRAY
    prod_p = doc.add_paragraph()
    prod_run = prod_p.add_run(f'Productos: {products}')
    prod_run.font.size = Pt(9)
    prod_run.bold = True
    prod_run.font.color.rgb = TEXT
    doc.add_paragraph()

def add_slide(doc, slide_num, slide_title, texto_imagen, frase=None, nota_visual=None, cta=None, marca=None):
    h = doc.add_heading(f'SLIDE {slide_num} — {slide_title}', level=2)
    for run in h.runs:
        run.font.color.rgb = PRIMARY_LIGHT
    
    p = doc.add_paragraph()
    ti = p.add_run(f'Texto imagen: ')
    ti.bold = True
    ti.font.size = Pt(9)
    ti.font.color.rgb = TEXT
    ti_val = p.add_run(f'"{texto_imagen}"')
    ti_val.font.size = Pt(9)
    ti_val.font.color.rgb = TEXT
    
    if frase:
        p2 = doc.add_paragraph()
        f = p2.add_run(f'Frase acompañante: ')
        f.bold = True
        f.font.size = Pt(9)
        f.font.color.rgb = TEXT
        f_val = p2.add_run(frase)
        f_val.font.size = Pt(9)
        f_val.font.color.rgb = TEXT
    
    if nota_visual:
        p3 = doc.add_paragraph()
        nv = p3.add_run(f'Nota visual: ')
        nv.bold = True
        nv.font.size = Pt(9)
        nv.font.color.rgb = GRAY
        nv_val = p3.add_run(nota_visual)
        nv_val.font.size = Pt(9)
        nv_val.font.color.rgb = GRAY
    
    if cta:
        p4 = doc.add_paragraph()
        cta_label = p4.add_run(f'CTA: ')
        cta_label.bold = True
        cta_label.font.size = Pt(9)
        cta_label.font.color.rgb = PRIMARY
        cta_val = p4.add_run(cta)
        cta_val.font.size = Pt(9)
        cta_val.font.color.rgb = PRIMARY
    
    if marca:
        p5 = doc.add_paragraph()
        m = p5.add_run(marca)
        m.bold = True
        m.font.size = Pt(9)
        m.font.color.rgb = PRIMARY
    
    doc.add_paragraph()

# ═══════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('COPYS S1 — JULIO 2026')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = DARK_BG

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Carruseles: #4 · #6 · #8 · #10')
run.font.size = Pt(14)
run.font.color.rgb = PRIMARY

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('Formato corto — listo para diseñadora')
run.font.size = Pt(11)
run.font.color.rgb = GRAY

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('Julio 2026 · SUPRICOM')
run.font.size = Pt(11)
run.font.color.rgb = GRAY

doc.add_page_break()

# ═══════════════════════════════════════════════════
# PIEZA #1 — BLUETTI: estaciones disponibles en SUPRICOM
# ═══════════════════════════════════════════════════
add_slide_title(doc, 1, 'BLUETTI: las estaciones de energia disponibles en SUPRICOM',
                'Carrusel (5 slides)', 'Showcase', 'BLUETTI (AC50P, AC70P, Elite 100 V2, Premium 200 V2, AC180P, AC200P L, Apex 300, Elite 30 V2)', 'No')

add_slide(doc, 1, 'Portada',
          'BLUETTI: las estaciones de energia disponibles en SUPRICOM',
          'Ocho modelos. Cada cliente tiene el suyo. Ya en Caracas y Valencia.',
          'Fondo azul degradado. Logo BLUETTI + SUPRICOM. Silueta de estacion de energia.')

add_slide(doc, 2, 'Linea portatil',
          'Para el dia a dia: compactas y ligeras',
          'BLUETTI AC50P y Elite 30 V2. Ideales para hogar, teletrabajo y llevar a cualquier lado.',
          'Icono mochila. AC50P y Elite 30 V2 lado a lado.')

add_slide(doc, 3, 'Linea media',
          'Para el comercio y la oficina: potencia balanceada',
          'BLUETTI AC70P, Elite 100 V2 y AC180P. Respaldo solido para POS, router y equipos de trabajo.',
          'Icono edificio. AC70P, Elite 100 V2, AC180P en fila.')

add_slide(doc, 4, 'Linea alta capacidad',
          'Para clinicas, talleres y demanda intensiva',
          'BLUETTI Premium 200 V2, AC200P L y Apex 300. Maxima capacidad, modo UPS, cero interrupciones.',
          'Icono cruz + herramientas. Premium 200 V2, AC200P L, Apex 300.')

add_slide(doc, 5, 'Cierre',
          'El portafolio BLUETTI completo, ya disponible.',
          'Ocho modelos para ocho perfiles de cliente. Pregunta por el que necesitas.',
          'SUPRICOM. Tu mayorista de confianza.',
          'Disponible en Caracas y Valencia.',
          'Todos los modelos BLUETTI en grid. Logo SUPRICOM.')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# PIEZA #4 — BLUETTI vs Planta Eléctrica
# ═══════════════════════════════════════════════════
add_slide_title(doc, 4, 'BLUETTI vs Planta Eléctrica: lo que tu cliente necesita saber',
                'Carrusel (5 slides)', 'Comparativa', 'BLUETTI (AC50P, AC70P, Premium 200 V2, Elite 100 V2)', 'No')

add_slide(doc, 1, 'Portada',
          'BLUETTI vs Planta Eléctrica: ¿qué conviene más?',
          'Ruido, mantenimiento, seguridad. Lo que necesitas saber para asesorar.',
          'Fondo azul degradado. Silueta BLUETTI izq., silueta generador der. Texto blanco.')

add_slide(doc, 2, 'Ruido',
          'Factor #1: Una funciona en silencio, la otra no',
          'BLUETTI es más silencioso que una biblioteca. Adentro de una oficina o casa, sin molestar.',
          'Icono ondas sonoras. Planta: rojo. BLUETTI: verde.')

add_slide(doc, 3, 'Mantenimiento',
          'Factor #2: Cero mantenimiento vs. cambios constantes',
          'Sin aceite, sin filtros, sin gasolina que se degrade. Enchufas y usas.',
          'Icono herramientas vs. icono enchufe.')

add_slide(doc, 4, 'Seguridad',
          'Factor #3: Cero emisiones, uso interior sin riesgo',
          'BLUETTI no produce monóxido. Seguro adentro de locales, oficinas y hogares.',
          'Señal peligro vs. hoja verde. Edificio en corte.')

add_slide(doc, 5, 'Cierre',
          'La respuesta que tu cliente necesita: BLUETTI',
          'Silencio, cero mantenimiento, uso interior. Tú decides según su perfil.',
          'SUPRICOM. Tu mayorista de confianza.',
          'Disponible en Caracas y Valencia.',
          'Logos BLUETTI, FORZA, SMARTBITT. Fondo blanco.')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# PIEZA #6 — Laptop segun perfil de cliente
# ═══════════════════════════════════════════════════
add_slide_title(doc, 6, 'Que laptop recomendar segun el perfil de tu cliente?',
                'Carrusel (6 slides)', 'Perfil', 'Acer · ASUS · Dell · HP · Lenovo', 'No')

add_slide(doc, 1, 'Portada',
          'Que laptop segun el perfil de tu cliente?',
          'Estudiante, oficinista, gerente, creativo. Cada uno con el modelo exacto.',
          'Fondo azul SUPRICOM. 4 iconos de perfil en esquinas.')

add_slide(doc, 2, 'Perfil #1 - Estudiante',
          'Perfil #1: Office, navegador, videollamadas',
          'Acer Notebook AGO15IN R5 16GB 512GB (NXJG6AL001). Sencilla y funcional para el dia a dia.',
          'Icono mochila. Laptop Acer en escritorio de estudio.')

add_slide(doc, 3, 'Perfil #2 - Oficinista',
          'Perfil #2: CRM, multitarea, videollamadas frecuentes',
          'ASUS VivoBook 15 X1502VA i7-13620H 16GB (90NB10T1-M01NH0). Potencia y fluidez.',
          'Icono oficina. Pantalla con hojas de calculo.')

add_slide(doc, 4, 'Perfil #3 - Gerente viajero',
          'Perfil #3: Bateria, portabilidad, resistencia profesional',
          'HP ProBook 4 G1i Core Ultra 7 255U 16" (BB3R1UT). Construida para movilidad.',
          'Icono maletin. Profesional con HP ProBook.')

add_slide(doc, 5, 'Perfil #4 - Creativo',
          'Perfil #4: Edicion, diseno, trabajo pesado',
          'ASUS TUF Gaming A15 Ryzen 7 RTX3050 (90NR0JF7-M00JA0). Rendimiento maximo.',
          'Icono monitor + paleta. Ambiente creativo con ASUS TUF.')

add_slide(doc, 6, 'Cierre',
          'Cada perfil tiene su laptop ideal.',
          'Portafolio completo: Acer, ASUS, Dell, HP, Lenovo. Pregunta por el SKU.',
          'SUPRICOM. Tu mayorista de confianza.',
          'Disponible en Caracas y Valencia.',
          'Grid logos Acer, ASUS, Dell, HP, Lenovo. Fondo blanco.')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# PIEZA #8 — 4 tipos de clientes que necesitan BLUETTI
# ═══════════════════════════════════════════════════
add_slide_title(doc, 8, '4 tipos de clientes que necesitan BLUETTI',
                'Carrusel (6 slides)', 'Perfil', 'BLUETTI (AC50P, AC70P, Premium 200 V2, Elite 100 V2)', 'No')

add_slide(doc, 1, 'Portada',
          '4 tipos de clientes que necesitan BLUETTI',
          'Cada uno con el modelo exacto que busca. Identifica y recomienda.',
          'Fondo azul. 4 iconos: casa, oficina, cruz, taller.')

add_slide(doc, 2, 'Cliente #1 — Hogar',
          'Cliente #1: Nevera, router, televisor. Lo esencial.',
          'BLUETTI AC50P: compacto, silencioso, horas de respaldo para el hogar.',
          'Icono casa. Familia con BLUETTI encendiendo luces y router.')

add_slide(doc, 3, 'Cliente #2 — Oficina / Comercio',
          'Cliente #2: POS, internet, cámaras. El negocio no para.',
          'BLUETTI AC70P: potencia media para mantener todo encendido sin interrupciones.',
          'Icono edificio. Mostrador con POS activo, afuera oscuro.')

add_slide(doc, 4, 'Cliente #3 — Clínica / Consultorio',
          'Cliente #3: Equipos médicos, medicamentos, historias digitales.',
          'BLUETTI Premium 200 V2: alta capacidad, modo UPS, cero riesgos.',
          'Icono cruz médica. Equipos funcionando sin interrupción.')

add_slide(doc, 5, 'Cliente #4 — Taller / Emprendimiento',
          'Cliente #4: Herramientas, diagnóstico, producción sin pausa.',
          'BLUETTI Elite 100 V2: portátil, potente, sin planta ruidosa.',
          'Icono herramientas. Taller con BLUETTI alimentando equipos.')

add_slide(doc, 6, 'Cierre',
          'Un cliente, una solución, un argumento.',
          'Vende con criterio. BLUETTI disponible en Caracas y Valencia.',
          'SUPRICOM. Tu mayorista de confianza.',
          'Disponible en Caracas y Valencia.',
          'Grid 4 iconos + logos BLUETTI y SUPRICOM. Fondo blanco.')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# PIEZA #10 — Inversión vs gasto: laptop con protección
# ═══════════════════════════════════════════════════
add_slide_title(doc, 10, 'Inversión vs gasto: laptop con protección',
                'Carrusel (5 slides)', 'Inversión', 'Acer · ASUS · Dell · HP · Lenovo · FORZA', 'No')

add_slide(doc, 1, 'Portada',
          'Laptop con UPS: ¿inversión o gasto?',
          'Dos escenarios. Lo único que cambia es si ofreces protección.',
          'Fondo azul. Laptop + UPS FORZA. Escudo de protección.')

add_slide(doc, 2, 'Sin protección',
          'Escenario A: Sin protección. El riesgo es real.',
          'Cortes y fluctuaciones dañan componentes. Garantía no cubre daños eléctricos.',
          'Laptop con rayo. Fondo rojo suave. Alerta.')

add_slide(doc, 3, 'Con protección',
          'Escenario B: Con UPS FORZA. La laptop está segura.',
          'Regulación AVR constante. Tiempo para guardar y apagar. El UPS se paga solo.',
          'Laptop + UPS FORZA. Fondo verde suave. Escudo.')

add_slide(doc, 4, 'Beneficio invisible',
          'Protección 24/7: incluso cuando no hay apagón.',
          'Fluctuaciones silentes dañan componentes día tras día. El UPS alarga la vida útil.',
          'Reloj con escudo. Línea de tiempo, laptop siempre protegida.')

add_slide(doc, 5, 'Cierre',
          'Vender protección es cuidar la inversión de tu cliente.',
          'Un UPS FORZA es la diferencia entre vender un equipo y una solución completa.',
          'SUPRICOM. Tu mayorista de confianza.',
          'Disponible en Caracas y Valencia.',
          'FORZA + SUPRICOM. Laptop con UPS. Fondo blanco.')

# ═══════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Documento generado Julio 2026 —')
run.font.color.rgb = GRAY
run.font.size = Pt(9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Copys listos para diseñadora • Formato corto')
run.font.color.rgb = GRAY
run.font.size = Pt(9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SUPRICOM. Tu mayorista de confianza.')
run.font.color.rgb = PRIMARY
run.bold = True

output = r'C:\Users\MDIGITAL01\Desktop\MARKETING OPENCODE\copys_s1_julio2026_v2.docx'
doc.save(output)
print(f'Documento guardado en: {output}')
